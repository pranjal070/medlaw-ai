import os
from pypdf import PdfReader
import docx

def extract_text_from_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also extract from tables in docx
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text).strip()
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")

def extract_document_text(file_path: str, filename: str) -> str:
    _, ext = os.path.splitext(filename.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".txt"]:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    else:
        # For images, we return empty string and let Gemini's multimodal API process it directly
        return ""
